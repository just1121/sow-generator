import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
import os
import io
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn  # Added qn here
from docx.shared import RGBColor
import logging
import inspect
import traceback
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pathlib import Path
import asyncio
from functools import partial
from streamlit import cache_data, cache_resource
from google.cloud import speech
import tempfile
import numpy as np
import base64
import json
import datetime
import time
import re
import streamlit.components.v1 as components
from streamlit_audio_recorder import st_audio_recorder
from pydub import AudioSegment
import subprocess
from google.cloud import storage

# Set page config as the first Streamlit command
st.set_page_config(layout="wide")

# Initialize section2_counter at the start of the script
section2_counter = 0

# # Add counter here, after imports but before any functions
# section2_counter = 0

# def load_test_data(scenario='default'):
#     """Load test data from JSON file"""
#     try:
#         current_dir = os.path.dirname(os.path.abspath(__file__))
#         test_data_path = os.path.join(current_dir, 'test_data.json')
#         with open(test_data_path, 'r') as f:
#             test_scenarios = json.load(f)
#             return test_scenarios.get(scenario, test_scenarios['default'])
#     except Exception as e:
#         st.error(f"Error loading test data: {e}")
#         return None

# def populate_test_data():
#     """Populate session state with test data"""
#     test_data = load_test_data()
#     if test_data:
#         for key, value in test_data.items():
#             st.session_state[key] = value
            
#         # Convert date strings to datetime objects
#         if 'effective_date' in st.session_state:
#             st.session_state.effective_date = datetime.datetime.strptime(st.session_state.effective_date, '%Y-%m-%d').date()
#         if 'master_terms_date' in st.session_state:
#             st.session_state.master_terms_date = datetime.datetime.strptime(st.session_state.master_terms_date, '%Y-%m-%d').date()
        
#         # Convert dates in deliverables
#         if 'deliverables' in st.session_state:
#             for key, deliverable in st.session_state.deliverables.items():
#                 if 'target_date' in deliverable:
#                     deliverable['target_date'] = datetime.datetime.strptime(deliverable['target_date'], '%Y-%m-%d').date()
#                 if 'milestones' in deliverable:
#                     for milestone in deliverable['milestones']:
#                         if 'due_date' in milestone:
#                             milestone['due_date'] = datetime.datetime.strptime(milestone['due_date'], '%Y-%m-%d').date()
        
#         st.success("Test data populated successfully!")
#         return True
#     return False

# # Initialize test mode in session state
# if 'test_mode' not in st.session_state:
#     st.session_state.test_mode = False

# # Add toggle in sidebar
# with st.sidebar:
#     previous_state = st.session_state.test_mode
#     st.session_state.test_mode = st.toggle("Test Mode", st.session_state.test_mode)
#     if st.session_state.test_mode and not previous_state:
#         populate_test_data()  # Use the existing function instead of duplicating code

# Define initialize_session_state function first
def initialize_session_state():
    # Only initialize if not in test mode
    if 'test_mode' not in st.session_state:
        st.session_state.test_mode = False

    # Only initialize these if they don't exist and we're not in test mode
    if not st.session_state.test_mode:
        if 'generated_content' not in st.session_state:
            st.session_state.generated_content = None
        if 'additional_statements' not in st.session_state:
            st.session_state.additional_statements = ""
        if 'input_method' not in st.session_state:
            st.session_state.input_method = {}
        if 'questions' not in st.session_state:
            st.session_state.questions = {
                'client': {"id": 0, "question": "Who is the client?", "type": "text", "answer": ""},
                
                'general_info': [
                    {"id": 1, "question": "Why were we called?", "type": "text", "answer": ""},
                    {"id": 2, "question": "What is the strategic purpose of this project?", "type": "text", "answer": ""},
                    {"id": 3, "question": "What is the tactical goal of this project?", "type": "text", "answer": ""},
                    {"id": 4, "question": "What is the desired end state?", "type": "text", "answer": ""},
                    {"id": 5, "question": "What is the scope of work proposed?", "type": "text", "answer": ""},
                    {"id": 6, "question": "What is the technical approach to this project?", "type": "text", "answer": ""},
                    {"id": 7, "question": "What equipment is required?", "type": "text", "answer": ""},
                    {"id": 8, "question": "Why is that equipment important and unique?", "type": "text", "answer": ""},
                ],
                
                'project_details': {
                    'main_questions': [
                        {"id": 1, "question": "Are there specific technical requirements or limits to meet?", "type": "radio", "answer": ""}
                    ],
                    'technical_details': [
                        {"id": 1, "question": "What is the technical approach to this project?", "type": "text", "answer": ""},
                        {"id": 2, "question": "What are the specific technical requirements?", "type": "text", "answer": ""},
                        {"id": 3, "question": "What are the technical limits?", "type": "text", "answer": ""},
                        {"id": 4, "question": "Are there any regulatory standards to comply with?", "type": "text", "answer": ""}
                    ]
                },
                
                'additional_details': [
                    {"id": 1, "question": "What are the potential risks and mitigation strategies?", "type": "text", "answer": ""},
                    {"id": 2, "question": "Are there intangible results expected? (expectations of client)", "type": "text", "answer": ""},
                    {"id": 3, "question": "Is coordination with other contractors required?", "type": "radio", "answer": "No"}
                ],
                'coordination_details': [
                    {"id": 1, "question": "Which contractors?", "type": "text", "answer": ""},
                    {"id": 2, "question": "What is the nature of coordination?", "type": "text", "answer": ""}
                ]
            }
        if 'effective_date' not in st.session_state:
            st.session_state.effective_date = datetime.date.today().isoformat()
        if 'master_terms_date' not in st.session_state:
            st.session_state.master_terms_date = datetime.date.today().isoformat()
        if 'client_address' not in st.session_state:
            st.session_state.client_address = ""
        if 'deliverables' not in st.session_state:  # Added this
            st.session_state.deliverables = {}
        if 'labor_costs' not in st.session_state:
            st.session_state.labor_costs = {
                'roles': [
                    {'role': 'Project Management', 'rate': 325.00, 'hours': 0},
                    {'role': 'Senior Wastewater Consultant', 'rate': 275.00, 'hours': 0},
                    {'role': 'Wastewater Consultant', 'rate': 225.00, 'hours': 0},
                    {'role': 'Senior Winemaker', 'rate': 250.00, 'hours': 0},
                    {'role': 'Winemaker', 'rate': 200.00, 'hours': 0},
                    {'role': 'Process Engineer', 'rate': 225.00, 'hours': 0},
                    {'role': 'Mechanical Engineer', 'rate': 225.00, 'hours': 0},
                    {'role': 'Fabrication Specialist', 'rate': 175.00, 'hours': 0},
                    {'role': 'Discipline Specialist', 'rate': 175.00, 'hours': 0},
                    {'role': 'Operation/Training Technician - ST', 'rate': 150.00, 'hours': 0},
                    {'role': 'Operation/Training Technician - OT', 'rate': 225.00, 'hours': 0},
                    {'role': 'AI Support', 'rate': 200.00, 'hours': 0},
                    {'role': 'Administration/Purchasing', 'rate': 135.00, 'hours': 0},
                    {'role': 'Schedule Administration', 'rate': 135.00, 'hours': 0},
                    {'role': 'Cost Administration', 'rate': 135.00, 'hours': 0}
                ],
                'expenses': {
                    'mileage_rate': 0.625,
                    'mileage': 0,
                    'truck_days': 0,
                    'truck_rate': 200.00,
                    'materials_cost': 0,
                    'materials_markup': 0.25
                },
                'total': 0
            }
        if 'expenses' not in st.session_state:  # Added this
            st.session_state.expenses = {
                'materials_cost': 0.0,
                'materials_markup': 0.25,    # Correct markup
                'mileage': 0,
                'mileage_rate': 0.625,      # Correct rate
                'truck_days': 0,
                'truck_rate': 200.00        # Correct rate
            }

# Add at the start of your app
try:
    client = speech.SpeechClient()
except Exception as e:
    st.error(f"Error creating Speech-to-Text client: {e}")

# Add near the top of your file, where other st.markdown calls are
st.markdown("""
    <style>
    #recordButton_audio_client::before {
        content: "Speak" !important;
    }
    #recordButton_audio_client[data-recording="true"]::before {
        content: "Stop" !important;
    }
    /* Hide clear recording buttons */
    button[kind="secondary"]:has(div:contains("Clear Recording")) {
        display: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

def load_environment():
    load_dotenv()
    if not os.getenv("GOOGLE_API_KEY"):
        st.error("Google API key not found. Please check your .env file.")
        st.stop()
    if not os.getenv("GOOGLE_APPLICATION_CREDENTIALS"):
        st.error("Google Application Credentials not found. Please set the GOOGLE_APPLICATION_CREDENTIALS environment variable.")
        st.stop()

def create_document(content, file_format):
    global section2_counter  # Move it here
    try:
        buffer = io.BytesIO()
        doc = Document()
        
        def apply_heading_style(paragraph):
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.font.bold = True
                
        def apply_body_style(paragraph):
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

        def format_table(table):
            """Apply consistent formatting to tables"""
            table.style = 'Table Grid'
            # Make header row bold
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            # Set column widths if needed
            table.autofit = False
            table.allow_autofit = False

        if content:
            paragraphs = [p for p in content.split('\n\n') if p.strip()]
            
            # Handle the first paragraph (document name) specially
            if paragraphs:
                first_para = doc.add_paragraph(paragraphs[0])
                first_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in first_para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(12)
            
            # Initialize flags before processing paragraphs
            in_section_2 = False
            in_section_3 = False
            in_section_5 = False
            
            for paragraph in paragraphs[1:]:
                # Skip content if we're in section 2 and see markdown
                if in_section_2 and any(marker in paragraph for marker in [
                    "Contractor will provide Deliverables under this SOW as described here:",
                    "**2. Description of Deliverables**",
                    "**Deliverable",
                    "| Milestone |",
                    "|-----------|",
                    "| Description |",
                    "| Target Date |"
                ]):
                    continue

                # Skip content if we're in section 5 and see markdown
                if in_section_5 and ('|' in paragraph or '**Deliverable' in paragraph):
                    continue

                if "5. Basis for Compensation" in paragraph:
                    in_section_5 = True  # Set flag when entering Section 5
                    p = doc.add_paragraph("5. Basis for Compensation")
                    apply_heading_style(p)
                    
                    # Add narrative
                    client_name = st.session_state.questions['client']["answer"].strip()
                    labor_cost = sum(details['total'] for deliverable in st.session_state.deliverables.values() 
                                  for details in deliverable.get('labor_costs', {}).values() 
                                  if isinstance(details, dict))
                    additional_expenses = (
                        st.session_state.expenses['mileage'] * st.session_state.expenses['mileage_rate'] +
                        st.session_state.expenses['truck_days'] * st.session_state.expenses['truck_rate'] +
                        st.session_state.expenses['materials_cost'] * (1 + st.session_state.expenses['materials_markup'])
                    )
                    total_cost = labor_cost + additional_expenses
                    
                    narrative = (f"The estimated cost for completion of this scope of work is ${total_cost:,.2f}. "
                               f"The tables below details the estimated efforts required. Material changes to the "
                               f"SOW will be agreed upon in writing and may constitute a change in basis for "
                               f"compensation increasing or decreasing accordingly.\n\n"
                               f"Efforts not explicitly listed in the table below are the responsibility of {client_name}. "
                               f"These efforts include, but are not limited to, onsite laboratory testing, offsite laboratory "
                               f"testing, and pilot system operation.")
                    
                    p = doc.add_paragraph(narrative)
                    apply_body_style(p)
                    
                    # Add deliverable labor cost tables
                    for i, (del_key, deliverable) in enumerate(st.session_state.deliverables.items(), 1):
                        if isinstance(deliverable.get('labor_costs'), dict):
                            # Add just one heading with the deliverable number and description
                            p = doc.add_paragraph(f"Deliverable {i}: {deliverable.get('description', '')}")
                            apply_heading_style(p)
                            
                            table = doc.add_table(rows=1, cols=5)
                            format_table(table)
                            header_cells = table.rows[0].cells
                            header_cells[0].paragraphs[0].add_run('Role').bold = True
                            header_cells[1].paragraphs[0].add_run('Description').bold = True
                            header_cells[2].paragraphs[0].add_run('Rate').bold = True
                            header_cells[3].paragraphs[0].add_run('Hours').bold = True
                            header_cells[4].paragraphs[0].add_run('Subtotal').bold = True
                            
                            has_labor_entries = False
                            for role, details in deliverable['labor_costs'].items():
                                if isinstance(details, dict) and details.get('hours', 0) > 0:
                                    has_labor_entries = True
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = role
                                    row_cells[1].text = details.get('description', '')  # Add description
                                    row_cells[2].text = f"${details['rate']:.2f}/hr"
                                    row_cells[3].text = f"{details['hours']:.2f}"
                                    row_cells[4].text = f"${details['total']:,.2f}"
                            
                            if has_labor_entries:
                                total_deliverable_cost = sum(details['total'] for details in deliverable['labor_costs'].values() if isinstance(details, dict))
                                doc.add_paragraph()  # Add space before total
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                p.add_run("Total Labor Cost for Deliverable: ")
                                p.add_run(f"${total_deliverable_cost:,.2f}").bold = True
                    
                    # Add Additional Expenses table
                    p = doc.add_paragraph("Additional Expenses")
                    apply_heading_style(p)
                    
                    table = doc.add_table(rows=1, cols=3)
                    format_table(table)
                    
                    # Set headers for Additional Expenses table
                    header_cells = table.rows[0].cells
                    header_cells[0].paragraphs[0].add_run('Expense Type').bold = True
                    header_cells[1].paragraphs[0].add_run('Details').bold = True
                    header_cells[2].paragraphs[0].add_run('Amount').bold = True
                    
                    # Add expense rows
                    expenses = st.session_state.expenses
                    
                    # Mileage row
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Mileage'
                    # Format mileage as whole number if it's a whole number, otherwise show decimals
                    mileage = expenses['mileage']
                    mileage_str = str(int(mileage)) if mileage.is_integer() else f"{mileage:.1f}"
                    row_cells[1].text = f"{mileage_str} miles @ ${expenses['mileage_rate']:.3f}/mile"
                    mileage_total = expenses['mileage'] * expenses['mileage_rate']
                    row_cells[2].text = f"${mileage_total:,.2f}"
                    
                    # Truck Days row
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Truck Days'
                    # Format truck days as whole number if it's a whole number, otherwise show decimals
                    truck_days = expenses['truck_days']
                    truck_days_str = str(int(truck_days)) if truck_days.is_integer() else f"{truck_days:.1f}"
                    row_cells[1].text = f"{truck_days_str} days @ ${expenses['truck_rate']:.2f}/day"
                    truck_total = expenses['truck_days'] * expenses['truck_rate']
                    row_cells[2].text = f"${truck_total:,.2f}"
                    
                    # Materials row
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Materials'
                    row_cells[1].text = f"Including {expenses['materials_markup']*100:.1f}% markup"
                    materials_total = expenses['materials_cost'] * (1 + expenses['materials_markup'])
                    row_cells[2].text = f"${materials_total:,.2f}"
    
                    # Add spacing after table
                    doc.add_paragraph()
                    
                    # Add rows for expenses
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Additional Costs'
                    row_cells[1].text = ''  # Empty details column
                    row_cells[2].text = f'${additional_expenses:,.2f}'  # Amount in last column
                    
                    # Add Project Totals table
                    p = doc.add_paragraph("Project Totals")
                    apply_heading_style(p)
                    
                    table = doc.add_table(rows=1, cols=2)
                    format_table(table)
                    header_cells = table.rows[0].cells
                    header_cells[0].paragraphs[0].add_run('Category').bold = True
                    header_cells[1].paragraphs[0].add_run('Amount').bold = True
                    
                    # Add rows for totals
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Total Labor Costs'
                    row_cells[1].text = f'${labor_cost:,.2f}'
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Total Additional Costs'
                    row_cells[1].text = f'${additional_expenses:,.2f}'
                    
                    row_cells = table.add_row().cells
                    p1 = row_cells[0].paragraphs[0]
                    p2 = row_cells[1].paragraphs[0]
                    p1.add_run('Total Project Cost').bold = True
                    p2.add_run(f'${total_cost:,.2f}').bold = True
                    
                    doc.add_paragraph()  # Add spacing after Section 5
                    continue

                # Skip any content between Sections 5 and 6
                if any(marker in paragraph for marker in [
                    "| Role |",
                    "Total Labor Cost for Deliverable",
                    "**Deliverable",
                    "|------|",
                    "|------",
                    "The estimated cost for completion of this scope of work",
                    "Efforts not explicitly listed in the table below",
                    "**Additional Expenses**",
                    "Project Totals"
                ]):
                    continue

                if "**Executive Summary**" in paragraph:
                    p = doc.add_paragraph("Executive Summary")
                    apply_heading_style(p)
                    continue

                elif "2. Description of Deliverables" in paragraph:
                    # Check processed flag immediately
                    if hasattr(doc, 'section2_processed'):
                        continue
                        
                    in_section_2 = True  # Set flag when entering Section 2
                    
                    # Add section header
                    p = doc.add_paragraph("2. Description of Deliverables")
                    apply_heading_style(p)
                    
                    # Add intro text
                    p = doc.add_paragraph("Contractor will provide Deliverables under this SOW as described here:")
                    apply_body_style(p)
                    
                    # Format each deliverable with milestones
                    for i, (del_key, deliverable) in enumerate(st.session_state.deliverables.items(), 1):
                        if deliverable.get('description', '').strip():
                            deliverable_text = f"Deliverable {i} {deliverable['description']}"
                            
                            # Add milestone information if exists
                            if deliverable.get('milestones') and len(deliverable['milestones']) > 0:
                                milestone_count = sum(1 for m in deliverable['milestones'] if m.get('description'))
                                
                                if milestone_count == 1:
                                    # Single milestone format
                                    milestone = next(m for m in deliverable['milestones'] if m.get('description'))
                                    deliverable_text += f" has one milestone, {milestone.get('description')}."
                                else:
                                    # Multiple milestones format
                                    deliverable_text += f" has {milestone_count} milestones"
                            
                            # Add period if not already present
                            if not deliverable_text.endswith('.'):
                                deliverable_text += '.'
                                
                            p = doc.add_paragraph(deliverable_text)
                            apply_body_style(p)

                # Add comprehensive skip conditions
                elif in_section_2 and any(marker in paragraph for marker in [
                    "Contractor will provide Deliverables under this SOW as described here",
                    "**2. Description of Deliverables**",
                    "**Deliverable",
                    "• Deliverable",
                    "- Deliverable",
                    "| Milestone |",
                    "|-----------|",
                    "| Description |",
                    "| Target Date |"
                ]):
                    continue

                elif "3. Work Schedule" in paragraph:
                    # Check processed flag immediately
                    if hasattr(doc, 'section3_processed'):
                        continue
                        
                    in_section_3 = True  # Set flag when entering Section 3
                    
                    # Add section header
                    p = doc.add_paragraph("3. Work Schedule")
                    apply_heading_style(p)

                    # Get completion date directly from session state
                    completion_date = st.session_state.get('expected_completion_date')
                    
                    try:
                        formatted_completion_date = completion_date.strftime('%B %d, %Y')
                    except Exception as e:
                        st.error(f"Error formatting completion date: {str(e)}")
                        formatted_completion_date = str(completion_date)

                    # Add the introductory text
                    client_name = st.session_state.questions['client']["answer"].strip()
                    intro_text = f"Contractor will conduct the Services and provide the Deliverables to {client_name} by {formatted_completion_date}. Specific deliverables and their projected timing are included below."
                    p = doc.add_paragraph(intro_text)
                    apply_body_style(p)

                    # Process each deliverable with bullet points and target dates
                    for i, (del_key, deliverable) in enumerate(st.session_state.deliverables.items(), 1):
                        if deliverable.get('description') and deliverable.get('target_date'):
                            p = doc.add_paragraph()
                            p.style = 'List Bullet'  # This gives us the bullet point (•) format
                            target_date = deliverable['target_date'].strftime('%B %d, %Y')
                            p.add_run(f"{deliverable.get('description')} by {target_date}")

                    # Mark section as processed
                    doc.section3_processed = True
                    continue

                # Skip any content that would create duplicate entries
                elif any(skip_text in paragraph for skip_text in [
                    "Contractor will conduct the Services and provide the Deliverables",
                    "Deliverable Schedule:",
                    "Project End Date:",
                    "- "  # This will catch all dash-formatted lines
                ]):
                    continue

                elif "4. Term of this SOW" in paragraph or "**4. Term of this SOW**" in paragraph:
                    p = doc.add_paragraph("4. Term of this SOW")
                    apply_heading_style(p)
                    continue

                elif "5. Basis for Compensation" in paragraph:
                    in_section_5 = True  # Set flag when entering Section 5
                    p = doc.add_paragraph("5. Basis for Compensation")
                    apply_heading_style(p)
                    
                    # Add narrative
                    client_name = st.session_state.questions['client']["answer"].strip()
                    labor_cost = sum(details['total'] for deliverable in st.session_state.deliverables.values() 
                                  for details in deliverable.get('labor_costs', {}).values() 
                                  if isinstance(details, dict))
                    additional_expenses = (
                        st.session_state.expenses['mileage'] * st.session_state.expenses['mileage_rate'] +
                        st.session_state.expenses['truck_days'] * st.session_state.expenses['truck_rate'] +
                        st.session_state.expenses['materials_cost'] * (1 + st.session_state.expenses['materials_markup'])
                    )
                    total_cost = labor_cost + additional_expenses
                    
                    narrative = (f"The estimated cost for completion of this scope of work is ${total_cost:,.2f}. "
                               f"The tables below details the estimated efforts required. Material changes to the "
                               f"SOW will be agreed upon in writing and may constitute a change in basis for "
                               f"compensation increasing or decreasing accordingly.\n\n"
                               f"Efforts not explicitly listed in the table below are the responsibility of {client_name}. "
                               f"These efforts include, but are not limited to, onsite laboratory testing, offsite laboratory "
                               f"testing, and pilot system operation.")
                    
                    p = doc.add_paragraph(narrative)
                    apply_body_style(p)
                    
                    # Add deliverable labor cost tables
                    for i, (del_key, deliverable) in enumerate(st.session_state.deliverables.items(), 1):
                        if isinstance(deliverable.get('labor_costs'), dict):
                            p = doc.add_paragraph(f"Deliverable {i}")
                            apply_heading_style(p)
                            
                            p = doc.add_paragraph(f"Description: {deliverable.get('description', '')}")
                            apply_body_style(p)
                            
                            table = doc.add_table(rows=1, cols=4)
                            format_table(table)
                            header_cells = table.rows[0].cells
                            header_cells[0].paragraphs[0].add_run('Role').bold = True
                            header_cells[1].paragraphs[0].add_run('Description').bold = True
                            header_cells[2].paragraphs[0].add_run('Rate').bold = True
                            header_cells[3].paragraphs[0].add_run('Hours').bold = True
                            header_cells[4].paragraphs[0].add_run('Subtotal').bold = True
                            
                            has_labor_entries = False
                            for role, details in deliverable['labor_costs'].items():
                                if isinstance(details, dict) and details.get('hours', 0) > 0:
                                    has_labor_entries = True
                                    row_cells = table.add_row().cells
                                    role_text = f"{role} - {details.get('description', '')}"
                                    row_cells[0].text = role_text
                                    row_cells[1].text = f"${details['rate']:.2f}/hr"
                                    row_cells[2].text = f"{details['hours']:.2f}"
                                    row_cells[3].text = f"${details['total']:,.2f}"
                            
                            if has_labor_entries:
                                total_deliverable_cost = sum(details['total'] for details in deliverable['labor_costs'].values() if isinstance(details, dict))
                                doc.add_paragraph(f"Total Cost for Deliverable: ${total_deliverable_cost:,.2f}")
                                doc.add_paragraph()
                    
                    # Add Additional Expenses table
                    p = doc.add_paragraph("Additional Expenses")
                    apply_heading_style(p)
                    
                    table = doc.add_table(rows=1, cols=2)
                    format_table(table)
                    header_cells = table.rows[0].cells
                    header_cells[0].paragraphs[0].add_run('Category').bold = True
                    header_cells[1].paragraphs[0].add_run('Amount').bold = True
                    
                    # Add rows for expenses
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Additional Costs'
                    row_cells[1].text = f'${additional_expenses:,.2f}'
                    
                    # Add Project Totals table
                    p = doc.add_paragraph("Project Totals")
                    apply_heading_style(p)
                    
                    table = doc.add_table(rows=1, cols=2)
                    format_table(table)
                    header_cells = table.rows[0].cells
                    header_cells[0].paragraphs[0].add_run('Category').bold = True
                    header_cells[1].paragraphs[0].add_run('Amount').bold = True
                    
                    # Add rows for totals
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Total Labor Costs'
                    row_cells[1].text = f'${labor_cost:,.2f}'
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Total Additional Costs'
                    row_cells[1].text = f'${additional_expenses:,.2f}'
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Total Project Cost'
                    row_cells[1].text = f'${total_cost:,.2f}'
                    
                    doc.add_paragraph()  # Add spacing after Section 5
                    continue

                elif "6. Title and Risk of Loss" in paragraph:
                    in_section_5 = False  # Reset flag
                    p = doc.add_paragraph("6. Title and Risk of Loss")
                    apply_heading_style(p)
                    continue

                elif "7. Additional Representations and Warranties" in paragraph:
                    p = doc.add_paragraph("7. Additional Representations and Warranties")
                    apply_heading_style(p)
                    
                    warranty_text = ("In addition to the representations and warranties set forth in the Agreement, "
                                    "Contractor represents and warrants to " + st.session_state.questions['client']['answer'] + 
                                    " that (i) neither its performance under the Agreement or this SOW nor any Deliverable "
                                    "(nor " + st.session_state.questions['client']['answer'] + "'s use thereof) will misappropriate, "
                                    "infringe, violate or interfere with the intellectual property or other right of any third party, "
                                    "(ii) it is not aware of, and has not received any notice of, any encroachment or infringement "
                                    "related to the Deliverables of any proprietary rights of any third party or in any violation "
                                    "of the rights of any third party, (iii) no Deliverable will contain a virus or other program "
                                    "or technology designed to disrupt, damage, interfere with or provide unauthorized access to "
                                    "any software, hardware or system, and (iv) it will not lose or corrupt " + 
                                    st.session_state.questions['client']['answer'] + "'s data (including, without limitation, third-party data).")
                    
                    p = doc.add_paragraph(warranty_text)
                    apply_body_style(p)
                    continue

                elif "8. Additional Terms" in paragraph:
                    p = doc.add_paragraph("8. Additional Terms")
                    apply_heading_style(p)
                    continue

                elif "9. List of" in paragraph:
                    p = doc.add_paragraph("9. List of attached SOW Schedules")
                    apply_heading_style(p)
                    continue

                elif '|' in paragraph and 'Milestone' in paragraph:
                    table = doc.add_table(rows=1, cols=3)
                    format_table(table)
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Milestone'
                    header_cells[1].text = 'Description'
                    header_cells[2].text = 'Target Date'
                    
                    # Extract milestone data from markdown table
                    rows = [row.strip() for row in paragraph.split('\n') if '|' in row and '-|-' not in row][1:]
                    for row in rows:
                        cells = [cell.strip() for cell in row.split('|')[1:-1]]
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(cells):
                            row_cells[i].text = cell.replace('**', '')
                    
                    doc.add_paragraph()  # Add spacing
                    continue

                else:
                    # Handle section headers (any numbered section)
                    if paragraph.startswith('**') and paragraph.endswith('**'):
                        # Remove markdown and check if it's a numbered section
                        clean_text = paragraph.strip('**')
                        if any(f"{i}." in clean_text for i in range(1, 10)):  # Handles sections 1-9
                            p = doc.add_paragraph(clean_text)
                            apply_heading_style(p)
                        else:
                            p = doc.add_paragraph(paragraph)
                            apply_body_style(p)
                    else:
                        p = doc.add_paragraph(paragraph)
                        apply_body_style(p)

            # Handle attached schedules
            if hasattr(st.session_state, 'attached_schedules') and st.session_state.attached_schedules:
                for file in st.session_state.attached_schedules:
                    try:
                        if file.name.lower().endswith('.docx'):
                            doc.add_page_break()
                            schedule_doc = Document(file)
                            for element in schedule_doc.element.body:
                                doc.element.body.append(element)
                        else:
                            doc.add_paragraph(f"Note: {file.name} is provided as a separate file")
                    except Exception as schedule_error:
                        st.warning(f"Could not process schedule {file.name}: {str(schedule_error)}")
                        continue

        buffer.flush()
        doc.save(buffer)
        buffer.seek(0)
        
        if buffer.getbuffer().nbytes > 10 * 1024 * 1024:
            st.warning("Document is quite large, may take longer to download")
            
        return buffer
        
    except Exception as e:
        st.error(f"Error creating document: {str(e)}")
        st.write("Error details:", traceback.format_exc())
        return None

@cache_data
def query_vectorstore(query, k=5):
    vectorstore = load_vectorstore()
    results = vectorstore.similarity_search(query, k=k)
    return results

@cache_resource
def load_vectorstore():
    try:
        # Use the same path construction as in initialize_vectorstore
        vectorstore_path = os.path.join(os.path.dirname(__file__), 'config', 'vectorstore.faiss')
        print(f"\nAttempting to load vectorstore from: {vectorstore_path}")
        print(f"Path exists: {os.path.exists(vectorstore_path)}")
        print(f"Directory contents: {os.listdir(os.path.dirname(vectorstore_path))}")
        
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        print("Created embeddings object")
        
        vectorstore = FAISS.load_local(vectorstore_path, embeddings, allow_dangerous_deserialization=True)
        print("Successfully loaded vectorstore")
        return vectorstore
        
    except Exception as e:
        print(f"Error loading vectorstore: {str(e)}")
        print(f"Full error: {traceback.format_exc()}")
        return None

def display_labor_costs():
    st.subheader("Labor Cost Estimation")
    
    # Labor Costs
    total_labor_cost = 0
    st.markdown("##### Labor Rates")
    
    def format_hours(hours):
        """Convert decimal hours to HH:MM format"""
        whole_hours = int(hours)
        minutes = int((hours - whole_hours) * 60)
        return f"{whole_hours}:{minutes:02d}"

    def parse_hours(time_str):
        """Convert HH:MM format to decimal hours"""
        try:
            if ':' in time_str:
                hours, minutes = map(int, time_str.split(':'))
                return hours + minutes/60
            return float(time_str)
        except:
            return 0.0
    
    for i, role in enumerate(st.session_state.labor_costs['roles']):
        col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
        
        with col1:
            st.text(role['role'])
        with col2:
            st.text(f"${role['rate']:.2f}/hr")
        with col3:
            hours = st.number_input(
                f"Hours", 
                value=float(role['hours']), 
                key=f"hours_{i}",
                min_value=0.0,
                step=0.25,  # Changed to 0.25 for quarter-hour increments
                format="%.2f",
                help="Enter time in quarter-hour increments (0.25 = 15 min, 0.5 = 30 min, 0.75 = 45 min)",
                label_visibility="collapsed"
            )
            displayed_time = format_hours(hours)
            st.text(f"({displayed_time})")  # Show HH:MM format below the input
            st.session_state.labor_costs['roles'][i]['hours'] = hours
        with col4:
            subtotal = role['rate'] * hours
            st.text(f"${subtotal:,.2f}")
            total_labor_cost += subtotal

    # Expenses
    st.markdown("##### Expenses")
    expenses = st.session_state.labor_costs['expenses']
    
    col1, col2 = st.columns(2)
    with col1:
        mileage = st.number_input(
            f"Mileage (${expenses['mileage_rate']:.3f}/mile)", 
            value=float(expenses['mileage']),
            min_value=0.0
        )
        truck_days = st.number_input(
            f"Truck Days (${expenses['truck_rate']:.2f}/day)", 
            value=int(expenses['truck_days']),
            min_value=0
        )
        materials_cost = st.number_input(
            f"Materials Cost (+ {expenses['materials_markup']*100}% markup)", 
            value=float(expenses['materials_cost']),
            min_value=0.0
        )
    
    with col2:
        mileage_total = mileage * expenses['mileage_rate']
        truck_total = truck_days * expenses['truck_rate']
        materials_total = materials_cost * (1 + expenses['materials_markup'])
        
        st.text(f"Mileage Total: ${mileage_total:,.2f}")
        st.text(f"Truck Total: ${truck_total:,.2f}")
        st.text(f"Materials Total: ${materials_total:,.2f}")

    # Update session state
    st.session_state.labor_costs['expenses']['mileage'] = mileage
    st.session_state.labor_costs['expenses']['truck_days'] = truck_days
    st.session_state.labor_costs['expenses']['materials_cost'] = materials_cost
    
    # Calculate grand total
    grand_total = total_labor_cost + mileage_total + truck_total + materials_total
    st.session_state.labor_costs['total'] = grand_total
    
    st.markdown("---")
    st.markdown(f"### Total Project Cost: ${grand_total:,.2f}")

def initialize_vectorstore():
    """Initialize the vectorstore with example documents"""
    try:
        print("Starting vectorstore initialization...")
        
        # Use relative path from the project root
        doc_path = os.path.join(os.path.dirname(__file__), 'config', 'documents')
        vectorstore_path = os.path.join(os.path.dirname(__file__), 'config', 'vectorstore.faiss')
        
        print(f"Loading documents from: {doc_path}")
        
        # Load PDF files only
        pdf_loader = DirectoryLoader(
            doc_path,
            glob="**/*.pdf",
            loader_cls=PyPDFLoader,
            recursive=True
        )
        documents = pdf_loader.load()
        print(f"Loaded {len(documents)} PDF documents")
        
        # Split documents
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        splits = text_splitter.split_documents(documents)
        print(f"Created {len(splits)} text chunks")
        
        # Use HuggingFace embeddings
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        
        # Create and save vectorstore
        vectorstore = FAISS.from_documents(splits, embeddings)
        vectorstore.save_local(vectorstore_path)
        
        print("Initialization complete!")
        return vectorstore
    except Exception as e:
        print(f"Error initializing vectorstore: {e}")
        return None

def get_answer(category, question_text):
    """Helper function to get answer for a specific question"""
    if category == 'client':
        return st.session_state.questions['client']["answer"]
    
    questions = st.session_state.questions.get(category, [])
    
    if isinstance(questions, list):
        for q in questions:
            if q["question"] == question_text:
                return q["answer"]
    elif isinstance(questions, dict):
        for section in questions.values():
            if isinstance(section, list):
                for q in section:
                    if q["question"] == question_text:
                        return q["answer"]
    return ""

def format_technical_requirements():
    """Format technical requirements if project has specific requirements"""
    tech_questions = st.session_state.questions['project_details']['technical_details']
    requirements = ""
    for q in tech_questions:
        if q["answer"]:
            requirements += f"- {q['question']}\n  {q['answer']}\n"
    return requirements.strip()

def format_labor_costs():
    """Format labor costs for all deliverables"""
    labor_lines = []
    total_cost = 0
    
    # Iterate through deliverables
    for i in range(st.session_state.get('deliverables_count', 1)):
        deliverable_key = f"deliverable_{i+1}"
        if deliverable_key in st.session_state.deliverables:
            deliverable = st.session_state.deliverables[deliverable_key]
            
            # Add deliverable header
            labor_lines.append(f"\nDeliverable {i+1}: {deliverable['description']}")
            
            # Add labor costs for this deliverable
            for role, details in deliverable.get('labor_costs', {}).items():
                if details['hours'] > 0:
                    cost = details['rate'] * details['hours']
                    labor_lines.append(f"  {role}: {details['hours']:.1f} hours at ${details['rate']:.2f}/hr (${cost:,.2f})")
                    total_cost += cost
    
    return f"""Customer shall feel the love from RWS as follows:

{chr(10).join(labor_lines)}

Total Project Cost: ${total_cost:,.2f}"""

def calculate_additional_costs():
    """Calculate all additional costs and store total in session state"""
    mileage_total = st.session_state.expenses['mileage'] * st.session_state.expenses['mileage_rate']
    truck_total = st.session_state.expenses['truck_days'] * st.session_state.expenses['truck_rate']
    materials_total = st.session_state.expenses['materials_cost'] * (1 + st.session_state.expenses['materials_markup'])
    st.session_state['total_additional_costs'] = mileage_total + truck_total + materials_total
    return mileage_total, truck_total, materials_total

# Add after your helper functions
async def check_vectorstore():
    try:
        # Search for chunks that start with EXECUTIVE SUMMARY
        results = await asyncio.to_thread(
            query_vectorstore, 
            "EXECUTIVE SUMMARY Grapevine Land Management seeks to develop"  # Using exact start of your example
        )
        
        for i, doc in enumerate(results, 1):
            # Clean up the display text
            content = doc.page_content
            # Fix word spacing issues
            content = content.replace("EXECUTIVESUMMARY", "EXECUTIVE SUMMARY")
            content = content.replace("  ", " ")
            # Split and rejoin to fix running-together words
            content = ' '.join([word for word in content.split() if word])
            
            st.write(f"\nDocument {i}:")
            st.write(content[:1000])
            st.write("-" * 80)
    except Exception as e:
        st.error(f"Error accessing vectorstore: {str(e)}")

def standardize_text(text):
    """Standardize text using dictionaries and regex patterns"""
    if not text:
        return text
        
    # Common spelling corrections
    spelling_corrections = {
        'recieved': 'received',
        'occured': 'occurred',
        # Add more as needed
    }
    
    # Date patterns
    date_patterns = {
        r'\b(\d{1,2})/(\d{1,2})/(\d{4})\b': lambda m: datetime.datetime.strptime(
            f"{m.group(1)}/{m.group(2)}/{m.group(3)}", 
            "%m/%d/%Y"
        ).strftime("%B %d, %Y"),
        r'\b([A-Za-z]{3,9})\s+(\d{1,2})\s*,?\s*(\d{4})\b': lambda m: datetime.datetime.strptime(
            f"{m.group(1)} {m.group(2)} {m.group(3)}", 
            "%B %d %Y"
        ).strftime("%B %d, %Y")
    }
    
    # Apply spelling corrections
    for wrong, right in spelling_corrections.items():
        text = re.sub(r'\b' + wrong + r'\b', right, text, flags=re.IGNORECASE)
    
    # Standardize dates
    for pattern, formatter in date_patterns.items():
        text = re.sub(pattern, formatter, text)
    
    # Fix common punctuation issues
    text = re.sub(r'\s+([.,!?])', r'\1', text)  # Remove space before punctuation
    text = re.sub(r'([.,!?])(?!["\'\s])', r'\1 ', text)  # Add space after punctuation
    
    return text

def standardize_user_content():
    """Standardize all user-entered content before SOW generation"""
    if 'deliverables' in st.session_state:
        for del_key, deliverable in st.session_state.deliverables.items():
            # Standardize deliverable description
            if 'description' in deliverable:
                deliverable['description'] = standardize_text(deliverable['description'])
            
            # Standardize labor cost descriptions
            if 'labor_costs' in deliverable:
                for role, details in deliverable['labor_costs'].items():
                    if isinstance(details, dict) and 'description' in details:
                        details['description'] = standardize_text(details['description'])

    # Standardize other user-entered fields from questions
    if 'questions' in st.session_state:
        for key, qa in st.session_state.questions.items():
            if isinstance(qa, dict) and 'answer' in qa and isinstance(qa['answer'], str):
                qa['answer'] = standardize_text(qa['answer'])

def generate_sow():  # no longer async
    with st.spinner("Generating SOW..."):
        try:
            # Standardize all user content first
            standardize_user_content() 

            # Then get all the user inputs (now standardized)
            client_name = st.session_state.questions['client']["answer"].strip()
            client_address = st.session_state.client_address
            effective_date = st.session_state.effective_date

            # Generate legal preamble
            legal_preamble = generate_legal_preamble(client_name, client_address, effective_date, st.session_state.master_terms_date)

            # Format deliverables in natural language
            deliverables_text = "\n\n**2. Description of Deliverables**\n\n"
            deliverables_text += "Contractor will provide Deliverables under this SOW as described here\n\n"
            
            for deliverable_index, (deliverable_key, deliverable) in enumerate(st.session_state.deliverables.items(), 1):
                if deliverable.get('description'):
                    # Start with deliverable description
                    deliverable_text = f"- Deliverable {deliverable_index} {deliverable.get('description')}"
                    
                    # Add milestone information if exists
                    if 'milestones' in deliverable and any(m.get('description') for m in deliverable['milestones']):
                        milestone_count = sum(1 for m in deliverable['milestones'] if m.get('description'))
                        
                        if milestone_count == 1:
                            # Single milestone format
                            milestone = next(m for m in deliverable['milestones'] if m.get('description'))
                            deliverable_text += f" has one milestone, {milestone.get('description')}."
                        else:
                            # Multiple milestones format
                            deliverable_text += f" has {milestone_count} milestones\n"
                            for milestone in deliverable['milestones']:
                                if milestone.get('description'):
                                    deliverable_text += f"  - {milestone.get('description')}"
                    
                    # Add period if not already present and no milestones were added
                    if not deliverable_text.endswith('.'):
                        deliverable_text += '.'
                        
                    deliverables_text += deliverable_text + "\n"

            prompt = f"""
            As a specialist for Recovered Water Solutions, please generate a Statement of Work (SOW).
            
            Executive Summary Generation Instructions:

1. Format Requirements:
- MUST contain exactly 3 paragraphs of flowing narrative
- Professional business story format
- Active voice with clear transitions

2. Paragraph Structure:
First Paragraph:
- Open with client's current situation and market context
- Identify critical challenges and capability gaps
- Create urgency for solution

Second Paragraph:
- Present RWS's expertise and unique capabilities
- Detail technical approach and methodology
- Highlight key equipment and processes
- Show understanding of critical challenges

Third Paragraph:
- Focus on strategic outcomes and long-term value
- Connect solution back to market opportunity
- Express confidence in project success
- End with client's enhanced future capabilities
            
            Using this information:
            {legal_preamble}

            Strategic Context:
            - Why Called: {get_answer('general_info', 'Why were we called?')}
            - Strategic Purpose: {get_answer('general_info', 'What is the strategic purpose of this project?')}
            - Tactical Goal: {get_answer('general_info', 'What is the tactical goal of this project?')}
            - Desired End State: {get_answer('general_info', 'What is the desired end state?')}

            Technical Approach:
            - Scope of Work: {get_answer('general_info', 'What is the scope of work proposed?')}
            - Technical Approach: {get_answer('general_info', 'What is the technical approach to this project?')}
            - Equipment Required: {get_answer('general_info', 'What equipment is required?')}
            - Equipment Importance: {get_answer('general_info', 'Why is that equipment important and unique?')}

            Project Structure:
            {format_technical_requirements() if st.session_state.get('tech_req') == "Yes" else ""}

            Implementation Requirements:
            - Training Needs: {get_answer('additional_details', 'What training is necessary for this project?')}
            - Reporting Requirements: {get_answer('additional_details', 'What reporting is necessary for this project?')}
            - Risk Management: {get_answer('additional_details', 'What are the potential risks and mitigation strategies?')}

            Additional context:
            {st.session_state.additional_statements}

            Your response should follow this exact format:
            ## Executive Summary
            [Generate 3 paragraphs following the structure above]

            ## 1. Description of Services
            [Detail the technical approach, methodologies, and specific services]
            """

            # Configure and generate with model
            genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(prompt).text  # Remove asyncio.to_thread

            if response:
                logging.info("Model response received")
                logging.info("=" * 50)
                logging.info("Raw model response:")
                logging.info(response)
                logging.info("=" * 50)

                # Split response into sections
                sections = response.split("##")
                
                # Clean the executive summary
                exec_summary = sections[1] if len(sections) > 1 else ""
                exec_summary = re.sub(r'.*?Executive\s+Summary\s*\n+', '', exec_summary, flags=re.IGNORECASE | re.MULTILINE)
                exec_summary = exec_summary.strip()
                
                # Get and clean services content
                services_content = sections[2] if len(sections) > 2 else ""
                services_content = re.sub(r'.*?Description\s+of\s+Services\s*\n+', '', services_content, flags=re.IGNORECASE | re.MULTILINE)
                services_content = re.sub(r'Technical\s+Approach\s*:', '\nTechnical Approach:', services_content)
                
                # Generate deliverables section from user inputs instead of model
                deliverables_content = deliverables_text
                
                # Standardize the services content
                services_content = standardize_text(services_content)
                
                # Generate remaining sections (without standardization)
                sections_3_to_9 = (
                    f"{generate_section_3()}\n\n"
                    f"{generate_section_4()}\n\n"
                    f"{generate_section_5_costs()}\n\n"
                    f"{generate_section_6(client_name)}\n\n"
                    f"{generate_section_7(client_name)}\n\n"
                    f"{generate_section_8()}\n\n"
                    f"{generate_section_9()}"
                )

                combined_content = (
                    f"[Document Name]\n\n"
                    f"{legal_preamble}\n\n"
                    f"**Executive Summary**\n\n"
                    f"{exec_summary}\n\n"
                    f"**1. Description of Services**\n\n"
                    f"{services_content}\n\n"
                    f"{deliverables_content}\n\n"
                    f"{sections_3_to_9}"
                )

                st.session_state.sow_result = {
                    'status': 'success',
                    'content': combined_content
                }
                logging.info("SOW generation successful")
            else:
                logging.error("No response from model")
                st.session_state.sow_result = {
                    'status': 'error',
                    'error': "No response received from model"
                }

        except Exception as e:
            logging.error(f"Error during SOW generation: {str(e)}", exc_info=True)
            st.session_state.sow_result = {
                'status': 'error',
                'error': str(e)
            }

def start_sow_generation():
    """Non-blocking function to start SOW generation"""
    if not st.session_state.get('sow_generation_started', False):
        st.session_state.sow_generation_started = True
        generate_sow()  # Removed asyncio.run() since function is no longer async

def create_entries_record():
    def format_table(table):
        """Apply consistent formatting to tables"""
        table.style = 'Table Grid'
        # Make header row bold
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        # Set column widths if needed
        table.autofit = False
        table.allow_autofit = False

    try:
        buffer = io.BytesIO()
        doc = Document()
        
        # Add title with error checking
        try:
            doc.add_heading('Statement of Work Question Responses', 0)
            doc.add_paragraph(f'Generated on: {datetime.datetime.now().strftime("%B %d, %Y %I:%M %p")}')
        except Exception as title_error:
            st.warning("Could not add title, continuing with content")
        
        # Client Info
        doc.add_heading('Client Information', level=1)
        doc.add_paragraph(f'Client: {st.session_state.questions["client"]["answer"]}')
        doc.add_paragraph(f'Address: {st.session_state.client_address}')
        
        # General Info (only if exists)
        if 'general_info' in st.session_state.questions:
            doc.add_heading('General Project Information', level=1)
            for question in st.session_state.questions['general_info']:
                doc.add_paragraph(f'{question["question"]}\nResponse: {question["answer"]}\n')
        
        # Project Details
        doc.add_heading('Project Details', level=1)

        # Add Deliverables section (only if deliverables exist)
        if 'deliverables' in st.session_state and st.session_state.deliverables:
            doc.add_heading('Deliverables, Milestones, and Associated Labor Costs', level=1)
            for key, deliverable in st.session_state.deliverables.items():
                if deliverable.get('description'):
                    doc.add_heading(f'Deliverable {key.split("_")[1]}', level=2)
                    doc.add_paragraph(f'Description: {deliverable["description"]}')
                    doc.add_paragraph(f'Target Completion Date: {deliverable["target_date"].strftime("%B %d, %Y")}')
                    
                    # Add Milestones
                    if deliverable.get('milestones'):
                        doc.add_heading('Milestones', level=3)
                        for i, milestone in enumerate(deliverable['milestones'], 1):
                            doc.add_paragraph(
                                f'Milestone {i}:\n'
                                f'Description: {milestone["description"]}\n'
                                f'Target Completion Date: {milestone["due_date"].strftime("%B %d, %Y")}'
                            )
                    
                    # Add Equipment and Additional Services
                    if deliverable.get('equipment_provided'):
                        doc.add_paragraph(f'Equipment and Materials Provided: {deliverable["equipment_provided"]}')
                    if deliverable.get('additional_services'):
                        doc.add_paragraph(f'Additional Services: {deliverable["additional_services"]}')
                    
                    # Only add labor costs table if there are labor costs
                    if deliverable.get('labor_costs'):
                        doc.add_heading('Labor Costs', level=3)
                        table = doc.add_table(rows=1, cols=5)
                        format_table(table)
                        header_cells = table.rows[0].cells
                        # Make headers bold
                        header_cells[0].paragraphs[0].add_run('Role').bold = True
                        header_cells[1].paragraphs[0].add_run('Description').bold = True
                        header_cells[2].paragraphs[0].add_run('Rate').bold = True
                        header_cells[3].paragraphs[0].add_run('Hours').bold = True
                        header_cells[4].paragraphs[0].add_run('Subtotal').bold = True
                        
                        # Only show roles with hours > 0
                        has_labor_entries = False
                        for role, details in deliverable['labor_costs'].items():
                            if details['hours'] > 0:
                                has_labor_entries = True
                                row_cells = table.add_row().cells
                                role_text = f"{role} - {details.get('description', '')}"
                                row_cells[0].text = role_text
                                row_cells[1].text = f"${details['rate']:.2f}/hr"
                                row_cells[2].text = f"{details['hours']:.2f}"
                                row_cells[3].text = f"${details['total']:,.2f}"
                        
                        if has_labor_entries:
                            total_deliverable_cost = sum(details['total'] for details in deliverable['labor_costs'].values())
                            doc.add_paragraph()  # Add space before total
                            p = doc.add_paragraph("Total Labor Cost for Deliverable: ")
                            p.add_run(f"${total_deliverable_cost:,.2f}").bold = True

        # Technical Requirements section
        tech_req = st.session_state.get('tech_req', 'No')
        doc.add_paragraph(f'Technical Requirements: {tech_req}')
        if tech_req == "Yes":
            doc.add_heading('Technical Requirements Details', level=2)
            if 'tech_1' in st.session_state:
                doc.add_paragraph(f'Technical Approach:\nResponse: {st.session_state.get("tech_1", "")}\n')
            if 'tech_2' in st.session_state:
                doc.add_paragraph(f'Specific Requirements:\nResponse: {st.session_state.get("tech_2", "")}\n')
            if 'tech_3' in st.session_state:
                doc.add_paragraph(f'Technical Limits:\nResponse: {st.session_state.get("tech_3", "")}\n')
            if 'tech_4' in st.session_state:
                doc.add_paragraph(f'Regulatory Standards:\nResponse: {st.session_state.get("tech_4", "")}\n')

        doc.add_paragraph()
        
        # Additional Statements
        doc.add_heading('Additional Statements', level=1)
        doc.add_paragraph(f'Response: {st.session_state.additional_statements}\n')

        # Calculate additional costs ONCE using the safer .get() method
        mileage_total = st.session_state.expenses.get('mileage', 0) * st.session_state.expenses.get('mileage_rate', 0)
        truck_total = st.session_state.expenses.get('truck_days', 0) * st.session_state.expenses.get('truck_rate', 0)
        materials_total = st.session_state.expenses.get('materials_cost', 0) * (1 + st.session_state.expenses.get('materials_markup', 0))
        total_additional_costs = mileage_total + truck_total + materials_total

        # Additional Costs section
        doc.add_heading(f'Additional Costs: ${total_additional_costs:,.2f}').bold = True
        
        # Display the costs
        doc.add_paragraph(f"Mileage Total: ${mileage_total:,.2f}")
        doc.add_paragraph(f"Truck Total: ${truck_total:,.2f}")
        doc.add_paragraph(f"Materials Total: ${materials_total:,.2f}")

        # Final Totals
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f"Total Additional Costs: ${total_additional_costs:,.2f}").bold = True
        p = doc.add_paragraph()
        p.add_run(f"Total Labor Costs: ${st.session_state['total_labor_cost']:,.2f}").bold = True
        p = doc.add_paragraph()
        p.add_run(f"Total Project Cost: ${(total_additional_costs + st.session_state['total_labor_cost']):,.2f}").bold = True

        # Save with retry logic
        max_retries = 3
        for attempt in range(max_retries):
            try:
                doc.save(buffer)
                buffer.seek(0)
                return buffer
            except Exception as save_error:
                if attempt == max_retries - 1:  # Last attempt
                    raise save_error
                time.sleep(0.5)  # Short delay before retry
                
    except Exception as e:
        st.error(f"Error creating entries record: {str(e)}")
        return None

def transcribe_audio(audio_content, sample_rate_hertz=16000):
    try:
        if isinstance(audio_content, dict):
            audio_data = audio_content.get('data', audio_content.get('bytes', []))
        else:
            audio_data = audio_content
            
        audio_bytes = bytearray(audio_data)
        
        with open('input.webm', 'wb') as f:
            f.write(audio_bytes)
            
        result = subprocess.run(
            ['ffmpeg', '-y',
             '-i', 'input.webm',
             '-acodec', 'pcm_s16le', 
             '-ar', str(sample_rate_hertz),
             '-ac', '1', 'output.wav'],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        
        with open('output.wav', 'rb') as f:
            wav_content = f.read()
            
        audio = speech.RecognitionAudio(content=wav_content)
        config = speech.RecognitionConfig(
            encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
            sample_rate_hertz=sample_rate_hertz,
            language_code="en-US"
        )
        
        response = client.recognize(config=config, audio=audio)
        transcript = response.results[0].alternatives[0].transcript
        
        return transcript
        
    except Exception as e:
        print(f"Error: {e}")
        return None

def get_audio_input(question, key):
    if key not in st.session_state:
        st.session_state[key] = ""
    
    # Add a transcription_complete flag
    transcription_key = f"transcription_complete_{key}"
    if transcription_key not in st.session_state:
        st.session_state[transcription_key] = False
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        text_value = st.text_area(question, 
                                key=f"text_{key}", 
                                value=st.session_state[key],
                                height=100)
        if text_value != st.session_state[key]:
            st.session_state[key] = text_value
    
    if st.session_state.input_method == 'Audio':
        with col2:
            status_placeholder = st.empty()
            prev_audio_key = f"prev_audio_{key}"
            if prev_audio_key not in st.session_state:
                st.session_state[prev_audio_key] = None
            
            audio_bytes = st_audio_recorder(key=f"audio_{key}")
            
            # Only process when audio_bytes changes from None to not None
            if (audio_bytes is not None and 
                st.session_state[prev_audio_key] is None and 
                not st.session_state[transcription_key]):
            
                status_placeholder.info("Processing audio...")
                transcription = transcribe_audio(audio_bytes)
                
                if transcription:
                    st.session_state[key] = transcription
                    st.session_state[transcription_key] = True
                    status_placeholder.success("Transcription complete!")
                    st.rerun()
                else:
                    status_placeholder.error("Transcription failed")
            
            # Reset transcription flag after rerun
            if st.session_state[transcription_key]:
                st.session_state[transcription_key] = False
            
            # Update previous audio state
            st.session_state[prev_audio_key] = audio_bytes
    
    return st.session_state[key]

def generate_legal_preamble(client_name, client_address, sow_date, master_terms_date):
    if not client_name:
        client_name = "[CLIENT NAME REQUIRED]"
    if not client_address:
        client_address = "[CLIENT ADDRESS REQUIRED]"
    if isinstance(sow_date, str):
        sow_date = datetime.date.fromisoformat(sow_date)
    if isinstance(master_terms_date, str):
        master_terms_date = datetime.date.fromisoformat(master_terms_date)
        
    sow_date_str = sow_date.strftime("%B %d, %Y")
    master_date_str = master_terms_date.strftime("%B %d, %Y")
    sow_year = sow_date.year

    preamble = f"""This statement of work ("SOW"), dated as of {sow_date_str} (the "SOW Effective Date") is by and between {client_name} ("Customer") with its principal place of business located at {client_address} and American Winesecrets, LLC, doing business as Recovered Water Solutions with its principal address at 1446 Industrial Avenue, Sebastopol, CA    95472, a California LLC ("Contractor" or "RWS"). This SOW and any accompanying exhibits, is incorporated into, forms a part of, and is in all respects subject to the terms of a Master Terms & Conditions Agreement (the "Agreement") dated {master_date_str}."""
    return preamble

def format_deliverables_section():
    """Generate Description of Deliverables section"""
    content = []
    
    # Check if deliverables exist
    if not st.session_state.deliverables:
        return "No deliverables have been defined.\n"
    
    # Add section header
    content.append("Contractor will provide Deliverables under this SOW as described here:\n")
    
    # Format each deliverable
    for i, (del_key, deliverable) in enumerate(st.session_state.deliverables.items(), 1):
        if deliverable.get('description', '').strip():  # Only include if description exists
            content.append(f"\n**Deliverable {i}: {deliverable['description']}**")
            
            # Add milestone table if milestones exist
            if deliverable.get('milestones') and len(deliverable['milestones']) > 0:
                content.append("\n| Milestone | Description | Target Date |")
                content.append("|-----------|-------------|-------------|")
                
                for j, milestone in enumerate(deliverable['milestones'], 1):
                    content.append(
                        f"| {j} | {milestone.get('description', '')} | "
                        f"{milestone['due_date'].strftime('%B %d, %Y')} |"
                    )
                content.append("")  # Add spacing after table
            
            # Add equipment and services if present
            if deliverable.get('equipment_provided'):
                content.append(f"\n**Equipment Provided:**\n{deliverable['equipment_provided']}")
            if deliverable.get('additional_services'):
                content.append(f"\n**Additional Services:**\n{deliverable['additional_services']}")
    
    return "\n".join(content)

def generate_section_3():
    """Generate Work Schedule section with improved formatting"""
    # Retrieve and strip the client name
    client_name = st.session_state.questions['client']["answer"].strip()
    
    # Attempt to retrieve the completion date from session state
    completion_date = st.session_state.get('expected_completion_date', None)

    # Check if completion_date is empty
    if not completion_date:
        st.error("Completion date is not provided.")
        return

    # Format completion date
    try:
        formatted_completion_date = completion_date.strftime('%B %d, %Y')
    except Exception as e:
        st.error(f"Error formatting completion date: {str(e)}")
        return

    # Build the section as a string
    section = "**3. Work Schedule**\n\n"
    
    section += (f"Contractor will conduct the Services and provide the Deliverables to {client_name} "
               f"by {formatted_completion_date}. Specific deliverables and their projected timing "
               f"are included below.\n\n")
    
    section += "**Deliverable Schedule:**\n\n"

    # Add deliverables to the schedule
    if st.session_state.deliverables:
        filled_count = 0
        for i, (del_key, deliverable) in enumerate(st.session_state.deliverables.items(), 1):
            description_filled = bool(deliverable.get('description', '').strip())
            # Maybe you want to check deliverable["target_date"] if it's mandatory
            if description_filled:
                filled_count += 1
                target_date = deliverable.get('target_date')
                if target_date:
                    date_str = target_date.strftime('%B %d, %Y')
                    section += f"- {deliverable['description']} by {date_str}\n"
                else:
                    section += f"- {deliverable['description']} (No date provided)\n"
        
        if filled_count == 0:
            section += "No work schedule has been defined.\n"
    else:
        section += "No work schedule has been defined.\n"

    return section

def generate_section_4():
    """Generate Section 4 with timeline data from session state"""
    
    # Get effective date directly from session state where it's already stored
    effective_date = st.session_state.effective_date
    
    # Get and standardize completion date
    completion_date = st.session_state.get('expected_completion_date', None)
    
    # Check if completion_date is empty
    if not completion_date:
        st.error("Completion date is not provided.")
        return

    # Format completion date
    try:
        formatted_completion_date = completion_date.strftime('%B %d, %Y')  # Convert to string
    except Exception as e:
        logging.warning(f"Could not standardize completion date: {str(e)}")
        return

    section_4 = "**4. Term of this SOW**\n\n"
    if effective_date and formatted_completion_date:
        section_4 += (f"This SOW shall be effective as of {effective_date.strftime('%B %d, %Y')} "
                     f"and shall remain in effect until the completion of all services and deliverables by {formatted_completion_date}.")
    else:
        section_4 += "[ERROR: Missing effective date or completion date]"
    
    return section_4

def generate_section_5_costs():
    """Generate Section 5: Basis for Compensation"""
    section = "**5. Basis for Compensation**\n\n"
    
    # Calculate totals
    total_labor_cost = sum(
        details.get('total', 0)
        for deliverable in st.session_state.deliverables.values()
        for details in deliverable.get('labor_costs', {}).values()
        if isinstance(details, dict)
    )
    
    # Calculate expenses totals
    expenses = st.session_state.expenses
    mileage_total = expenses['mileage'] * expenses['mileage_rate']
    truck_total = expenses['truck_days'] * expenses['truck_rate']
    materials_total = expenses['materials_cost'] * (1 + expenses['materials_markup'])
    total_additional = mileage_total + truck_total + materials_total
    
    # Add narrative
    client_name = st.session_state.questions['client']["answer"].strip()
    section += (f"The estimated cost for completion of this scope of work is ${total_labor_cost + total_additional:,.2f}. "
                f"The tables below details the estimated efforts required. Material changes to the "
                f"SOW will be agreed upon in writing and may constitute a change in basis for "
                f"compensation increasing or decreasing accordingly.\n\n"
                f"Efforts not explicitly listed in the table below are the responsibility of {client_name}. "
                f"These efforts include, but are not limited to, onsite laboratory testing, offsite laboratory "
                f"testing, and pilot system operation.\n\n")
    
    # Add deliverables and labor costs in table format
    for i, (del_key, deliverable) in enumerate(st.session_state.deliverables.items(), 1):
        if isinstance(deliverable.get('labor_costs'), dict):
            section += f"**Deliverable {i}: {deliverable.get('description', '')}**\n\n"
            
            # Add labor costs table
            section += "| Role | Description | Rate | Hours | Subtotal |\n"
            section += "|------|-------------|------|-------|----------|\n"
            
            for role, details in deliverable['labor_costs'].items():
                if isinstance(details, dict) and details.get('hours', 0) > 0:
                    section += (f"| {role} | {details.get('description', '')} | "
                              f"${details['rate']:.2f}/hr | {details['hours']:.2f} | "
                              f"${details['total']:,.2f} |\n")
            
            # Add deliverable total
            del_total = sum(details['total'] for details in deliverable['labor_costs'].values() 
                          if isinstance(details, dict))
            section += f"\n**Total Labor Cost for Deliverable: ${del_total:,.2f}**\n\n"
    
    # Add expenses section with table
    section += "\n**Additional Expenses**\n\n"
    section += "| Expense Type | Details | Amount |\n"
    section += "|--------------|----------|--------|\n"
    section += f"| Mileage | {expenses['mileage']} miles @ ${expenses['mileage_rate']:.3f}/mile | ${mileage_total:,.2f} |\n"
    section += f"| Truck Days | {expenses['truck_days']} days @ ${expenses['truck_rate']:.2f}/day | ${truck_total:,.2f} |\n"
    section += f"| Materials | Including {expenses['materials_markup']*100}% markup | ${materials_total:,.2f} |\n"
    
    # Add totals section
    section += "\n**Project Totals**\n\n"
    section += "| Category | Amount |\n"
    section += "|----------|--------|\n"
    section += f"| Total Additional Costs | ${total_additional:,.2f} |\n"
    section += f"| Total Labor Costs | ${total_labor_cost:,.2f} |\n"
    section += f"| **Total Project Cost** | **${total_labor_cost + total_additional:,.2f}** |\n"
    
    return section

def generate_section_6(client_name):
    if hasattr(st.session_state, 'title_terms') and st.session_state.title_terms:
        terms = st.session_state.title_terms.replace('[Client Name]', client_name)
        return f"\n**6. Title and Risk of Loss**\n\n{terms}"
    return f"\n**6. Title and Risk of Loss**\n\nN/A"

def generate_section_7(client_name):
    section = (
        "**7. Additional Representations and Warranties.**\n\n In addition to the representations and warranties set forth in the Agreement, "
        f"Contractor represents and warrants to {client_name} that (i) neither its performance under the Agreement or this SOW nor any "
        f"Deliverable (nor {client_name}'s use thereof) will misappropriate, infringe, violate or interfere with the intellectual "
        f"property or other right of any third party, (ii) it is not aware of, and has not received any notice of, any encroachment "
        f"or infringement related to the Deliverables of any proprietary rights of any third party or in any violation of the rights "
        f"of any third party, (iii) no Deliverable will contain a virus or other program or technology designed to disrupt, damage, "
        f"interfere with or provide unauthorized access to any software, hardware or system, and (iv) it will not lose or corrupt "
        f"{client_name}'s data (including, without limitation, third-party data)."
    ).replace("\n", " ").replace("  ", " ")
    
    return section

def generate_section_8():
    if hasattr(st.session_state, 'additional_terms') and st.session_state.additional_terms:
        return f"\n**8. Additional Terms**\n\n{st.session_state.additional_terms}"
    return f"\n**8. Additional Terms**\n\nNone."

def generate_section_9():
    if hasattr(st.session_state, 'attached_schedules') and st.session_state.attached_schedules:
        schedule_list = "\n• ".join(file.name for file in st.session_state.attached_schedules)
        return f"\n**9. List of attached SOW Schedules**\n\n {schedule_list}"
    return f"\n**9. List of attached SOW Schedules**\n\nNone"

def generate_section_5():
    """Generate Additional Terms or Payment Terms section, etc."""
    section = "**5. Payment Terms**\n\n"
    
    # Suppose we want to list deliverables that have cost info
    # (This is just an example – adapt as needed)
    if st.session_state.deliverables:
        valid_deliverables = 0
        for i, (del_key, deliverable) in enumerate(st.session_state.deliverables.items(), 1):
            description_filled = bool(deliverable.get('description', '').strip())
            if description_filled:
                valid_deliverables += 1
                # Suppose we print labor_costs if present
                labor_costs = deliverable.get('labor_costs', {})
                # ... format them, etc.
                section += f"- {deliverable['description']} (Cost details: {labor_costs})\n"
        
        if valid_deliverables == 0:
            section += "No payment terms have been defined.\n"
    else:
        section += "No payment terms have been defined.\n"

    return section

# ALTERNATIVE EXECUTIVE SUMMARY GENERATION - COMMENTED OUT FOR FUTURE REFERENCE
# async def generate_executive_summary(user_input, additional_statements):
#     prompt = f"""
#     Generate a compelling Executive Summary that tells the project's story. The summary should:
#     
#     Style and Tone:
#     - Create narrative flow with a clear beginning, middle, and resolution
#     - Build subtle tension by highlighting challenges/risks
#     - Show how your expertise provides the solution
#     - Maintain professional tone while engaging the reader
#     - Use active voice and confident language
#     
#     Key Elements to Weave Together:
#     - Current situation and critical challenges
#     - Stakes and implications for the client
#     - Your unique approach and expertise
#     - Implementation strategy
#     - Value proposition and expected outcomes
#     
#     Structure:
#     - Opening: Hook reader with context and challenge
#     - Middle: Present solution and expertise
#     - Close: Highlight benefits and confidence in success
#     
#     Using this input:
#     Project Details: {user_input}
#     Additional Context: {additional_statements}
#     
#     Create a flowing 2-3 paragraph narrative that builds trust while showcasing your understanding and capability.
#     Focus on telling a compelling story that demonstrates value and expertise.
#     """
#
#     # Generate using GenAI
#     response = await generate_content(prompt)
#     return response

def main():
    # Define CSS for styling
    st.markdown(
        """
        <style>
        .custom-divider {
            border-top: 4px solid #007BFF;  /* Blue color */
            margin: 15px 0;                 /* Space above and below */
        }
        /* Target Streamlit's specific input elements */
        textarea {
            background-color: #e6f7ff !important;
        }
        .stTextInput > div > div > input {
            background-color: #e6f7ff !important;
        }
        /* Target date input elements */
        .stDateInput > div > div > input {
            background-color: #e6f7ff !important;
        }
        /* Target number/selectbox input - slightly darker blue */
        .stSelectbox > div > div > div {
            background-color: #cce9ff !important;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Change this part near the start of main()
    if test_bucket_access():
        pass
    else:
        st.error("Could not connect to cloud storage bucket")

    initialize_session_state()
    load_environment()

    # Header and Logo
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        image_path = os.path.join(current_dir, "static", "RecWater-Solutions-Logo-1.png")
        if os.path.exists(image_path):
            st.image(image_path, width=200)
        else:
            st.error(f"Logo not found. Please ensure 'RecWater-Solutions-Logo-1.png' is in the 'static' folder.")

    with col2:
        st.markdown("<h1 style='text-align: center;'>Statement of Work Generator</h1>", unsafe_allow_html=True)

    st.markdown("---")
    
    st.session_state.input_method = st.radio("Choose input method:", ['Text', 'Audio'])

    # Client Information
    st.subheader("Client Information")
    st.session_state.questions['client']["answer"] = get_audio_input("Who is the client?", "client")
    
    col1, col2 = st.columns([2, 2])
    with col1:
        st.session_state.client_address = st.text_input(
            "What is the client's address?",
            value=st.session_state.get("client_address", "")
        )
        
        # Effective Date Input
        default_date = datetime.date.today()
        try:
            if isinstance(st.session_state.effective_date, str):
                default_date = datetime.date.fromisoformat(st.session_state.effective_date)
            elif isinstance(st.session_state.effective_date, datetime.date):
                default_date = st.session_state.effective_date
        except ValueError:
            pass
        st.session_state.effective_date = st.date_input(
            "What is the effective date of this Statement of Work?",
            value=default_date
        )
        
        # NEW: expected Completion Date Input
        default_completion_date = datetime.date.today()  # You can adjust the default if desired
        st.session_state.expected_completion_date = st.date_input(
            "What is the expected completion date for the project covered in this Statement of Work?",
            value=default_completion_date
        )
        
        # Master Terms Date Input
        default_master_date = datetime.date.today()
        try:
            if isinstance(st.session_state.master_terms_date, str):
                default_master_date = datetime.date.fromisoformat(st.session_state.master_terms_date)
            elif isinstance(st.session_state.master_terms_date, datetime.date):
                default_master_date = st.session_state.master_terms_date
        except ValueError:
            pass
        st.session_state.master_terms_date = st.date_input(
            "What is the effective date of the Master Terms & Conditions Agreement?",
            value=default_master_date
        )

    # General Project Information
    st.subheader("General Project Information")
    for question in st.session_state.questions['general_info']:
        question["answer"] = get_audio_input(question["question"], f"general_{question['id']}")
    
    # Deliverables section - correct way
    st.markdown("**How many deliverables does this project involve?**")
    col1, col2 = st.columns([1, 3])  # Create two columns with 1:3 ratio
    with col1:
        num_deliverables = st.selectbox(
        label="Number of deliverables",
        options=list(range(1, 10)),
        key="deliverables_count",
            label_visibility="collapsed"  # This hides the label but keeps it accessible
        )

    # Initialize deliverables in session state if not present
    if 'deliverables' not in st.session_state:
        st.session_state.deliverables = {}

    # Add a variable to keep track of the total labor cost for all deliverables
    if 'total_labor_cost' not in st.session_state:
        st.session_state['total_labor_cost'] = 0.0

    # Initialize the session state to store total labor costs across deliverables
    total_labor_cost = 0.0

    # Create sections for each deliverable
    for i in range(num_deliverables):
        section_class = "deliverable-section-even" if i % 2 == 0 else "deliverable-section-odd"
        st.markdown(f"<div class='{section_class}'>", unsafe_allow_html=True)
        
        st.markdown(f"### Deliverable {i + 1}")
        
        deliverable_key = f"deliverable_{i+1}"
        
        # Initialize this deliverable if it doesn't exist
        if deliverable_key not in st.session_state.deliverables:
            st.session_state.deliverables[deliverable_key] = {
                'description': '',
                'equipment_provided': '',
                'additional_services': '',
                'labor_costs': {},
                'target_date': datetime.date.today(),
                'milestones': []
            }
        
        # Main description
        # Main description (full width)
        description = st.text_area(
            f"Name or description of deliverable {i + 1}:",
            help="Describe the primary consultation or service\n" +
                 "Example: 'Onsite consultation for protocol development'",
            value=st.session_state.deliverables[deliverable_key]['description'],
            key=f"desc_{deliverable_key}"
        )

        # Target completion date (narrow)
        col1, col2 = st.columns([1, 3])
        with col1:
            target_date = st.date_input(
                f"Target completion date for deliverable {i + 1}:",
                value=st.session_state.deliverables[deliverable_key]['target_date'],
                key=f"target_date_{deliverable_key}"
            )

        # Milestones section
        st.markdown(f"**Milestones for Deliverable {i + 1}**")
        
        # Get current milestones from session state
        current_milestones = st.session_state.deliverables[deliverable_key].get('milestones', [])
        
        # Create a unique key for the milestone count selectbox
        milestone_count_key = f"num_milestones_{deliverable_key}"
        
        # Initialize the milestone count in session state if it doesn't exist
        if milestone_count_key not in st.session_state:
            st.session_state[milestone_count_key] = len(current_milestones)
        
        col1, col2 = st.columns([1, 3])
        with col1:
            num_milestones = st.selectbox(
                label="Number of milestones (leave at 0 if none for deliverable)",
                options=list(range(0, 10)),
                key=milestone_count_key
            )
        
        # Only adjust milestone list if the count has changed
        if num_milestones != len(current_milestones):
            # Preserve existing milestone data
            new_milestones = current_milestones.copy()
            
            # Add new milestones if needed
            while len(new_milestones) < num_milestones:
                new_milestones.append({
                    'description': '',
                    'due_date': target_date if len(new_milestones) == num_milestones - 1 else default_date
                })
            
            # Remove excess milestones if needed
            while len(new_milestones) > num_milestones:
                new_milestones.pop()
            
            # Update session state
            st.session_state.deliverables[deliverable_key]['milestones'] = new_milestones
            current_milestones = new_milestones

        # Display milestone inputs
        for j, milestone in enumerate(current_milestones):
            st.markdown(f"*Milestone {j + 1}*")
            
            # Create unique keys for milestone inputs
            desc_key = f"milestone_desc_{deliverable_key}_{j}"
            date_key = f"milestone_date_{deliverable_key}_{j}"
            
            # Milestone description
            milestone_desc = st.text_input(
                f"Description for milestone {j + 1}",
                value=milestone.get('description', ''),
                key=desc_key
            )

            # Milestone due date
            col1, col2 = st.columns([1, 3])
            with col1:
                if j == num_milestones - 1:
                    milestone_date = target_date
                    st.write(f"Target completion date for Deliverable {i + 1} and Milestone {j + 1}: {target_date.strftime('%B %d, %Y')}")
                else:
                    default_milestone_date = min(
                        milestone.get('due_date', default_date),
                        target_date
                    )
                    milestone_date = st.date_input(
                        f"Target completion date for milestone {j + 1}:",
                        value=default_milestone_date,
                        key=date_key,
                        max_value=target_date
                    )

            # Update milestone in session state
            st.session_state.deliverables[deliverable_key]['milestones'][j] = {
                'description': milestone_desc,
                'due_date': milestone_date
            }

        # Equipment and materials (after milestones)
        equipment = st.text_area(
            "Equipment and materials provided:",
            help="List any equipment or materials included\n" +
                 "Example: 'RO1110 system, wine vessels, system chiller'",
            value=st.session_state.deliverables[deliverable_key].get('equipment_provided', ''),
            key=f"equipment_{deliverable_key}"
        )

        # Additional services
        additional = st.text_area(
            "Additional services:",
            help="Any additional services or opportunities\n" +
                 "Example: 'Offering samples to potential customers'",
            value=st.session_state.deliverables[deliverable_key].get('additional_services', ''),
            key=f"additional_{deliverable_key}"
        )

        # Update session state
        st.session_state.deliverables[deliverable_key].update({
            'description': description,
            'equipment_provided': equipment,
            'additional_services': additional,
            'target_date': target_date,
            'milestones': current_milestones
        })

        # Labor Categories expander for this deliverable
        with st.expander(f"Labor Categories - Deliverable {i + 1}"):
            labor_categories = {
                "Project Management": {"rate": 325.00, "hours": 0},
                "Senior Wastewater Consultant": {"rate": 275.00, "hours": 0}, 
                "Wastewater Consultant": {"rate": 225.00, "hours": 0},
                "Senior Winemaker": {"rate": 250.00, "hours": 0},
                "Winemaker": {"rate": 200.00, "hours": 0},
                "Process Engineer": {"rate": 225.00, "hours": 0},
                "Mechanical Engineer": {"rate": 225.00, "hours": 0},
                "Fabrication Specialist": {"rate": 175.00, "hours": 0},
                "Discipline Specialist": {"rate": 175.00, "hours": 0},
                "Operation/Training Technician - ST": {"rate": 150.00, "hours": 0},
                "Operation/Training Technician - OT": {"rate": 225.00, "hours": 0},
                "AI Support": {"rate": 200.00, "hours": 0},
                "Administration/Purchasing": {"rate": 135.00, "hours": 0},
                "Schedule Administration": {"rate": 135.00, "hours": 0},
                "Cost Administration": {"rate": 135.00, "hours": 0}
            }
            
            # Initialize total cost for this deliverable
            total_deliverable_cost = 0.0
            
            for role, details in labor_categories.items():
                col1, col2, col3, col4 = st.columns([2, 1, 1, 1])
                
                with col1:
                    st.text(role)
                    # Create a unique key for tracking hours
                    hours_key = f"hours_{deliverable_key}_{role.replace(' ', '_')}"
                    
                    # Get existing hours from session state
                    existing_hours = st.session_state.deliverables.get(deliverable_key, {}).get('labor_costs', {}).get(role, {}).get('hours', 0)
                    
                    with col3:
                        hours = st.number_input(
                            f"Hours",
                            value=float(existing_hours),
                            key=hours_key,
                            min_value=0.0,
                            step=0.25,
                            format="%.2f"
                        )
                    
                    # Calculate total before displaying
                    details['total'] = details['rate'] * hours
                    
                    with col4:
                        st.text(f"${details['rate']:.2f}/hr")
                        st.text(f"${details['total']:,.2f}")
                    
                    # Show description field immediately if hours exist in session state
                    if hours > 0:
                        desc_key = f"work_desc_{deliverable_key}_{role.replace(' ', '_')}"
                        existing_desc = st.session_state.deliverables.get(deliverable_key, {}).get('labor_costs', {}).get(role, {}).get('description', '')
                        work_description = st.text_area(
                            "Description of Work",
                            value=existing_desc,
                            key=desc_key,
                            placeholder=f"Enter description for {role} work...",
                            label_visibility="collapsed"
                        )
                        
                        # Update session state immediately
                        if deliverable_key not in st.session_state.deliverables:
                            st.session_state.deliverables[deliverable_key] = {'labor_costs': {}}
                        if 'labor_costs' not in st.session_state.deliverables[deliverable_key]:
                            st.session_state.deliverables[deliverable_key]['labor_costs'] = {}
                        
                        st.session_state.deliverables[deliverable_key]['labor_costs'][role] = {
                            'hours': hours,
                            'rate': details['rate'],
                            'total': details['total'],
                            'description': work_description
                        }
                
                with col4:
                    st.text(f"${details['rate']:.2f}/hr")
                    
                    with col4:
                        st.text(f"${details['total']:,.2f}")
                    
                    # Update session state for this role
                    if deliverable_key not in st.session_state.deliverables:
                        st.session_state.deliverables[deliverable_key] = {'labor_costs': {}}
                    
                    # Only store description if hours > 0
                    if hours > 0:
                        st.session_state.deliverables[deliverable_key]['labor_costs'][role] = {
                            'hours': hours,
                            'rate': details['rate'],
                            'total': details['total'],
                            'description': work_description
                        }
                    else:
                        st.session_state.deliverables[deliverable_key]['labor_costs'][role] = {
                            'hours': hours,
                            'rate': details['rate'],
                            'total': details['total']
                        }
                
                # Update the global total labor cost in session state
            st.session_state['total_labor_cost'] += total_deliverable_cost

                # Show individual total for the deliverable
            st.markdown(f"**Total Cost for Deliverable {i + 1}: ${total_deliverable_cost:,.2f}**")

        st.markdown("</div>", unsafe_allow_html=True)

            # Add separator between deliverables
        if i < num_deliverables - 1:
            st.markdown('<hr class="custom-divider">', unsafe_allow_html=True)

    st.markdown("### Technical Requirements")
    tech_req = st.radio(
        "Are there specific technical requirements or limits to meet?",
        options=["Yes", "No"],
        index=1,
        key="tech_req"
    )
    
    if tech_req == "Yes":
        for question in st.session_state.questions['project_details']['technical_details']:
            question["answer"] = get_audio_input(question["question"], f"tech_{question['id']}")

    # Additional Project Details
    st.subheader("Additional Project Details")
    for question in st.session_state.questions['additional_details']:
        if question["type"] == "radio":
            question["answer"] = st.radio(
                question["question"],
                ["Yes", "No"],
                key=f"additional_{question['id']}",
                index=0 if st.session_state.get(f"additional_{question['id']}") == "Yes" else 1
            )
            if question["answer"] == "Yes":
                for coord_q in st.session_state.questions['coordination_details']:
                    coord_q["answer"] = get_audio_input(coord_q["question"], f"coord_{coord_q['id']}")
        else:
            question["answer"] = get_audio_input(question["question"], f"additional_{question['id']}")

    # Additional Statements (moved up)
    st.markdown("### Additional Statements")
    st.session_state.additional_statements = get_audio_input(
        "Enter any additional statements or context that should be included:",
        "additional_statements"
    )

    # Additional Costs
    st.markdown("---")
    st.subheader("Additional Costs")
    
    # Reset total labor cost before calculating
    st.session_state['total_labor_cost'] = 0.0
    
    # Calculate current total from deliverables
    for i in range(num_deliverables):
        deliverable_key = f"deliverable_{i+1}"
        if deliverable_key in st.session_state.deliverables:
            total_deliverable_cost = sum(
                details.get('total', 0) 
                for details in st.session_state.deliverables[deliverable_key].get('labor_costs', {}).values()
            )   
            st.session_state['total_labor_cost'] += total_deliverable_cost
            
    # Initialize expenses with correct rates and types
    if 'expenses' not in st.session_state:
        st.session_state.expenses = {
            'mileage_rate': 0.625,      # float
            'mileage': 0.0,             # float
            'truck_rate': 200.00,       # float
            'truck_days': 0.0,          # float
            'materials_cost': 0.0,       # float
            'materials_markup': 0.25     # float
        }
    
    col1, col2 = st.columns(2)
    with col1:
        mileage = st.number_input(
            f"Mileage ($0.625/mile)", 
            value=float(st.session_state.expenses['mileage']),
            min_value=0.0,
            step=1.0
        )
        st.session_state.expenses['mileage'] = float(mileage)
        
        truck_days = st.number_input(
            f"Truck Days ($200.00/day)", 
            value=float(st.session_state.expenses['truck_days']),
            min_value=0.0,
            step=1.0
        )
        st.session_state.expenses['truck_days'] = float(truck_days)
        
        materials_cost = st.number_input(
            f"Materials Cost (+ 25% markup)", 
            value=float(st.session_state.expenses['materials_cost']),
            min_value=0.0,
            step=1.0
        )
        st.session_state.expenses['materials_cost'] = float(materials_cost)

    # Totals column
    with col2:
        # Calculate totals
        mileage_total = st.session_state.expenses['mileage'] * st.session_state.expenses['mileage_rate']
        truck_total = st.session_state.expenses['truck_days'] * st.session_state.expenses['truck_rate']
        materials_total = st.session_state.expenses['materials_cost'] * (1 + st.session_state.expenses['materials_markup'])
        
        # Display totals with consistent formatting
        st.text(f"Mileage Total: ${mileage_total:.2f}")
        st.text(f"Truck Total: ${truck_total:.2f}")
        st.text(f"Materials Total: ${materials_total:.2f}")
    
    # Calculate total additional costs
    total_additional_costs = mileage_total + truck_total + materials_total
    
    # Define labor roles
    labor_roles = [
        "Project Management",
        "Senior Wastewater Consultant",
        "Wastewater Consultant",
        "Senior Winemaker",
        "Winemaker",
        "Process Engineer",
        "Mechanical Engineer",
        "Fabrication Specialist",
        "Discipline Specialist",
        "Operation/Training Technician - ST",
        "Operation/Training Technician - OT",
        "AI Support",
        "Administration/Purchasing",
        "Schedule Administration",
        "Cost Administration"
    ]
    
    # Get total labor costs from session state
    total_labor = sum(st.session_state[f"labor_{role}"]["total"] 
                     for role in labor_roles 
                     if f"labor_{role}" in st.session_state)
    
    # Calculate grand total before displaying
    grand_total = total_additional_costs + total_labor
    
    # Display all totals
    st.markdown("---")
    st.markdown(f"**Total Additional Costs: ${total_additional_costs:,.2f}**")
    st.markdown(f"**Total Labor Costs: ${st.session_state['total_labor_cost']:,.2f}**")  # Use the same value here
    st.markdown(f"**Total Project Cost: ${(total_additional_costs + st.session_state['total_labor_cost']):,.2f}**")

    # Title and Risk of Loss Conditional
    st.markdown("---")
    st.subheader("Title and Risk of Loss")
    has_title_terms = st.radio("Any change to Title and Risk of Loss?", ["No", "Yes"], key="has_title_terms")
    if has_title_terms == "Yes":
        title_terms = st.text_area("Please enter Title and Risk of Loss terms:", key="title_terms_text")
        if title_terms:
            st.session_state.title_terms = title_terms

    # Schedule Attachments
    st.markdown("---")
    st.subheader("Schedule Attachments")
    st.markdown("Please attach any schedule files, Word docs only (DOCX format):")

    col1, _ = st.columns([2, 2])  # Creates two columns but only uses the first one
    with col1:
        uploaded_files = st.file_uploader("Upload Schedules", 
                                        accept_multiple_files=True,
                                        type=['docx'])

    if uploaded_files:
        st.session_state.attached_schedules = uploaded_files

    # Single button for Record of Entries
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("Create Entries Record"):
            try:
                entries_doc = create_entries_record()
                if entries_doc is not None:
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    st.download_button(
                        label="Download Entries Record",
                        data=entries_doc,
                        file_name=f"sow_entries_record_{timestamp}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="Download a record of all entries",
                        key=f"entries_download_{timestamp}"
                    )
                    # Clean up
                    entries_doc.close()
            except Exception as e:
                st.error(f"Error creating entries record: {str(e)}")

    # Generate SOW Button
    st.markdown("---")
    if st.button("Generate SOW", key="generate_sow_button"):
        start_sow_generation()

    if st.session_state.get('sow_generation_started'):
        if 'sow_result' in st.session_state:
            if st.session_state.sow_result['status'] == 'success':
                st.success("SOW Generated Successfully!")
                st.session_state.generated_content = st.session_state.sow_result['content']
                st.session_state.sow_generation_started = False
            elif st.session_state.sow_result['status'] == 'error':
                st.error(f"Error generating SOW: {st.session_state.sow_result['error']}")
                st.session_state.sow_generation_started = False
        else:
            st.info("Generating SOW... Please wait...")

    # Display Generated Content
    if st.session_state.get('generated_content'):
        st.subheader("Generated Statement of Work:")
        st.info(st.session_state.generated_content)
        st.markdown("---")

        try:
            document = create_document(st.session_state.generated_content, "DOCX")
            entries_document = create_entries_record()
            
            if document is not None and entries_document is not None:
                timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                client_name = st.session_state.questions['client']["answer"].strip().replace(" ", "_")
                
                # Create filenames
                sow_filename = f"sow_{client_name}_{timestamp}.docx"
                entries_filename = f"entries_{client_name}_{timestamp}.docx"
                
                # Only save files to Cloud Storage once
                if not st.session_state.get("sow_uploaded", False):
                    save_to_gcloud_bucket(document, sow_filename)
                    save_to_gcloud_bucket(entries_document, entries_filename)
                    st.session_state.sow_uploaded = True

                # Provide download button to download the SOW locally only
                st.download_button(
                    label="Download Statement of Work",
                    data=document,
                    file_name=sow_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Download Statement of Work document",
                    key=f"sow_download_{int(time.time())}"
                )
        except Exception as e:
            st.error(f"Error processing documents: {str(e)}")

    if "audio_data" not in st.session_state:
        st.session_state.audio_data = None

    # Create a placeholder for the transcription
    transcription_placeholder = st.empty()

    # Handle the audio data when received
    if st.session_state.audio_data:
        transcription = handle_audio_data(st.session_state.audio_data)
        if transcription:
            transcription_placeholder.write(f"Transcription: {transcription}")
        st.session_state.audio_data = None  # Clear the audio data

def save_to_gcloud_bucket(file_buffer, filename, bucket_name="docxdownloads"):
    """Save a file to Google Cloud Storage bucket"""
    try:
        # Remove the "Attempting" message
        os.environ["GOOGLE_CLOUD_PROJECT"] = "recovered-water-solutions"
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(filename)
        
        file_buffer.seek(0)
        blob.upload_from_file(file_buffer)
        # Remove the success message
        return True
    except Exception as e:
        st.error(f"Error saving to Cloud Storage: {str(e)}")
        return False

def test_bucket_access():
    try:
        storage_client = storage.Client()
        bucket = storage_client.bucket("docxdownloads")
        blobs = list(bucket.list_blobs())
        # Remove the st.write message
        return True
    except Exception as e:
        st.error(f"Bucket access error: {str(e)}")
        return False

def add_deliverables_form():
    with st.form("deliverables_form"):
        st.write("Add Project Deliverables")
        
        if 'deliverables' not in st.session_state:
            st.session_state.deliverables = []
            
        deliverable_name = st.text_input("Deliverable Name")
        deliverable_date = st.date_input("Expected Completion Date")
        
        if st.form_submit_button("Add Deliverable"):
            st.session_state.deliverables.append({
                "name": deliverable_name,
                "date": deliverable_date.strftime('%B %d, %Y')
            })
            st.success("Deliverable added!")

    # Display current deliverables
    if st.session_state.deliverables:
        st.write("Current Deliverables:")
        for d in st.session_state.deliverables:
            st.write(f"- {d['name']}: {d['date']}")

if __name__ == "__main__":
    import sys
    # Check if running directly with Python
    if not sys.argv[0].endswith('streamlit'):
        if len(sys.argv) > 1 and sys.argv[1] == "--init-vectorstore":
            print("Initializing vectorstore...")
            initialize_vectorstore()
            sys.exit(0)
    
    # Normal Streamlit operation
    port = int(os.environ.get("PORT", 8080))
    main()