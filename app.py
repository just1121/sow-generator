from flask import Flask, render_template, request, jsonify, send_file
from google.cloud import speech
from docx import Document
import tempfile

app = Flask(__name__)

# Questions structure (in memory only)
QUESTIONS = {
    "client": {"question": "Who is the client?"},
    "general_info": [
        {"question": "Why were we called?"},
        {"question": "Where is the project location?"}
    ]
}

@app.route('/')
def index():
    return render_template('index.html', questions=QUESTIONS)

@app.route('/transcribe', methods=['POST'])
def transcribe():
    try:
        audio_file = request.files['audio']
        
        # Google Speech-to-Text
        client = speech.SpeechClient()
        audio_content = audio_file.read()
        audio = speech.RecognitionAudio(content=audio_content)
        config = speech.RecognitionConfig(
            encoding=speech.RecognitionConfig.AudioEncoding.WEBM_OPUS,
            language_code="en-US",
            enable_automatic_punctuation=True
        )
        
        response = client.recognize(config=config, audio=audio)
        transcription = " ".join(result.alternatives[0].transcript for result in response.results)
        
        return jsonify({'transcription': transcription})
    except Exception as e:
        print(f"Transcription error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/generate_sow', methods=['POST'])
def generate_sow():
    try:
        # Get all answers from the form
        answers = request.json
        
        # Create a simple document
        doc = Document()
        doc.add_heading('Statement of Work', 0)
        
        # Add client info
        doc.add_heading('Client Information', level=1)
        doc.add_paragraph(answers['client'])
        
        # Add general info
        doc.add_heading('Project Details', level=1)
        for i, answer in enumerate(answers['general_info']):
            doc.add_paragraph(f"{QUESTIONS['general_info'][i]['question']}: {answer}")
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name='Statement_of_Work.docx'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)